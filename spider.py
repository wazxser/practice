import urllib.request
from bs4 import BeautifulSoup
import chardet
import requests
import docx
from docx import Document
from docx.shared import Inches
import re
import time


file = open('suanfatimu.txt', 'w')
document = Document()
for i in range(392):
    s = ""
    page = requests.get('https://www.nowcoder.com/ta/review-ml/review?page=' + str(i+1))
    soup = BeautifulSoup(page.text, 'lxml')
    imgs = soup.findAll('img', attrs={"src": re.compile(r'https.*images')})
    for j in range(len(imgs)):
        img = imgs[j]
        link = img.get('src')
        with open('img.jpg', 'wb') as tu:
            response = requests.get(link).content
            tu.write(response)
            time.sleep(1)
        document.add_picture('img.jpg')
    # print(soup.select('div[class="design-answer-box"]')[0])
    s += str(i+1) + '、'
    s += soup.select('div[class="final-question"]')[0].text
    s += '\n'
    s += soup.select('div[class="design-answer-box"]')[0].text
    s += '\n'
    document.add_paragraph(s)

# file.write(s
document.save('suanfatimu.doc')
